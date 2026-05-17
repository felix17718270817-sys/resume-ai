# -*- coding: utf-8 -*-
# AI简历优化助手 - Streamlit 网页应用

import streamlit as st
import io
import os
import base64
import re
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from openai import OpenAI

# ==================== 页面基础配置 ====================
# 设置页面标题、图标和布局
st.set_page_config(
    page_title="AI简历优化助手",
    page_icon="📄",
    layout="wide",
)

# ==================== 视频背景（Base64 内嵌，路径无关）====================
video_path = os.path.join(os.path.dirname(__file__), "assets", "planet_remix_scene.webm")
with open(video_path, "rb") as f:
    video_base64 = base64.b64encode(f.read()).decode("utf-8")

st.markdown(f"""
<video autoplay muted loop playsinline preload="auto" id="bg-video">
    <source src="data:video/webm;base64,{video_base64}" type="video/webm">
</video>
""", unsafe_allow_html=True)

# ==================== 自定义样式 ====================
st.markdown("""
<style>
    /* ============================================================
       Design System — Light Glass AI SaaS
       Color Tokens:
         card-bg:         rgba(255, 255, 255, 0.88)  → 白色毛玻璃
         border-card:     rgba(148, 163, 184, 0.25)
         border-focus:    rgba(99, 102, 241, 0.45)
         text-primary:    #1e293b
         text-secondary:  #475569
         text-tertiary:   #64748b
         accent:          #2563eb → #6366f1 → #7c3aed → #a855f7
    ============================================================ */

    /* ----- 字体 ----- */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
    * { font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }

    /* Streamlit chrome: 隐藏默认元素 + 全部容器透明 */
    [data-testid="stHeader"], [data-testid="stToolbar"], #MainMenu, footer { display: none !important; }
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stAppViewContainer"] > div,
    [data-testid="stVerticalBlock"], [data-testid="stVerticalBlockBorderWrapper"] {
        background: transparent !important;
    }
    [data-testid="stVerticalBlockBorderWrapper"] { border: none !important; box-shadow: none !important; padding: 0 !important; margin: 0 !important; }
    [data-testid="stVerticalBlock"] { border: none !important; box-shadow: none !important; }
    .block-container { padding: 1rem 0 3rem !important; max-width: 1360px !important; position: relative; z-index: 1; }

    /* ----- 视频背景（固定全屏）----- */
    #bg-video {
        position: fixed;
        top: 0;
        left: 0;
        width: 100vw;
        height: 100vh;
        object-fit: cover;
        z-index: -10;
        pointer-events: none;
    }

    /* ----- 标题文字样式（无背景卡片）----- */
    .hero-title {
        text-align: center;
        font-size: 3.6rem;
        font-weight: 800;
        letter-spacing: -0.04em;
        background: linear-gradient(135deg, #2563eb 0%, #6366f1 30%, #7c3aed 60%, #a855f7 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 0.5rem;
        filter: drop-shadow(0 2px 4px rgba(37,99,235,0.25));
        text-shadow: none;
    }
    .hero-accent {
        width: 56px;
        height: 4px;
        background: linear-gradient(90deg, #2563eb, #7c3aed, #a855f7);
        border-radius: 2px;
        margin: 0.6rem auto 0.9rem auto;
        box-shadow: 0 0 10px rgba(99,102,241,0.25);
    }
    .hero-subtitle {
        text-align: center;
        font-size: 1.05rem;
        font-weight: 500;
        color: #475569;
        max-width: 500px;
        margin: 0 auto 1.75rem auto;
        line-height: 1.6;
    }

    /* ----- 统一白色玻璃卡片（输入面板 & 结果模块共用）----- */
    .glass-card, .module-card {
        background: rgba(255, 255, 255, 0.88);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        border: 1px solid rgba(148, 163, 184, 0.25);
        border-radius: 20px;
        padding: 1.75rem;
        margin-bottom: 1.25rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.04), 0 4px 20px rgba(0,0,0,0.06);
        transition: border-color 0.3s ease, box-shadow 0.3s ease;
    }
    .glass-card:hover, .module-card:hover {
        border-color: rgba(99, 102, 241, 0.30);
        box-shadow: 0 1px 3px rgba(0,0,0,0.06), 0 6px 28px rgba(99,102,241,0.08);
    }

    /* ----- 输入框 / 文本域：白色卡片 ----- */
    div[data-testid="stTextArea"] textarea {
        background: rgba(255, 255, 255, 0.92) !important;
        border: 1px solid rgba(148, 163, 184, 0.25) !important;
        border-radius: 14px !important;
        padding: 0.85rem 1rem !important;
        font-size: 0.92rem !important;
        line-height: 1.7 !important;
        color: #1e293b !important;
        caret-color: #6366f1 !important;
        transition: border-color 0.25s ease, box-shadow 0.25s ease, background 0.25s ease !important;
    }
    div[data-testid="stTextArea"] textarea:hover {
        border-color: rgba(99, 102, 241, 0.25) !important;
        background: rgba(255, 255, 255, 0.98) !important;
    }
    div[data-testid="stTextArea"] textarea:focus {
        background: rgba(255, 255, 255, 0.98) !important;
        border-color: rgba(99, 102, 241, 0.45) !important;
        box-shadow: 0 0 0 4px rgba(99,102,241,0.08), 0 0 20px rgba(99,102,241,0.06) !important;
        outline: none !important;
    }
    div[data-testid="stTextArea"] textarea::placeholder {
        color: #64748b !important;
        font-weight: 400 !important;
        opacity: 1 !important;
    }
    /* 卡片内 textarea label 隐藏 */
    .glass-card div[data-testid="stTextArea"] label {
        display: none;
    }

    /* ----- 文件上传：外层卡片风格 + 浅色主题色修正（不碰 pointer-events/z-index/transform/position）----- */
    [data-testid="stFileUploader"] {
        background: transparent !important;
    }
    [data-testid="stFileUploader"] section {
        background: rgba(255, 255, 255, 0.92) !important;
        border: 1px solid rgba(148, 163, 184, 0.25) !important;
        border-radius: 14px !important;
        padding: 1rem !important;
    }
    /* 仅修正暗色主题泄露导致的按钮/文字颜色，不改布局与交互 */
    [data-testid="stFileUploader"] button {
        background: rgba(255, 255, 255, 0.92) !important;
        color: #1e293b !important;
        border-color: rgba(148, 163, 184, 0.25) !important;
    }
    [data-testid="stFileUploader"] span {
        color: #475569 !important;
    }
    [data-testid="stFileUploader"] small {
        color: #64748b !important;
    }
    /* st.text / st.code 等原始文本块 */
    pre, code {
        color: #1e293b !important;
        background: rgba(255, 255, 255, 0.6) !important;
        border-radius: 8px;
        border: 1px solid rgba(148, 163, 184, 0.15);
    }

    /* ----- 主按钮：蓝→紫渐变 ----- */
    .stButton > button {
        background: linear-gradient(135deg, #2563eb 0%, #6366f1 35%, #7c3aed 70%, #a855f7 100%);
        color: #ffffff;
        border: none;
        border-radius: 14px;
        padding: 0.9rem 3rem;
        font-size: 1.05rem;
        font-weight: 600;
        letter-spacing: 0.02em;
        box-shadow: 0 2px 8px rgba(37,99,235,0.15), 0 0 24px rgba(99,102,241,0.15);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 16px rgba(37,99,235,0.25), 0 0 36px rgba(99,102,241,0.25);
    }
    .stButton > button:active {
        transform: translateY(0);
    }

    /* ----- 下载按钮：白色 outline 风格 ----- */
    .stDownloadButton > button {
        background: rgba(255, 255, 255, 0.80);
        color: #1e293b;
        border: 1px solid rgba(148, 163, 184, 0.25);
        border-radius: 12px;
        padding: 0.6rem 2rem;
        font-size: 0.9rem;
        font-weight: 600;
        transition: all 0.25s ease;
    }
    .stDownloadButton > button:hover {
        border-color: rgba(99, 102, 241, 0.40);
        color: #4f46e5;
        background: rgba(255, 255, 255, 0.95);
        box-shadow: 0 0 20px rgba(99,102,241,0.08);
        transform: translateY(-1px);
    }

    /* ----- 分隔线 ----- */
    hr {
        margin: 1.25rem 0;
        border: none;
        height: 1px;
        background: rgba(148, 163, 184, 0.18);
    }

    /* ----- 结果大标题 ----- */
    .module-card-icon { font-size: 1.4rem; margin-right: 0.4rem; }
    .module-card-title { font-size: 1.1rem; font-weight: 700; color: #1e293b; margin-bottom: 1rem;
        padding-bottom: 0.75rem; border-bottom: 1px solid rgba(148, 163, 184, 0.15); }
    .module-card-body { font-size: 0.95rem; color: #1e293b; line-height: 1.8; }
    .module-card-body p { margin-bottom: 0.5rem; }
    .module-card-body ul { padding-left: 1.25rem; }
    .module-card-body li { margin-bottom: 0.35rem; }
    .module-card-body strong, .module-card-body h2, .module-card-body h3 { color: #0f172a; }

    .results-heading { font-size: 1.5rem; font-weight: 700; color: #1e293b; margin-bottom: 1.25rem; }

    /* ATS 分析 UI */
    .ats-score-bar { height: 12px; border-radius: 6px; background: rgba(148,163,184,0.15); margin: 0.75rem 0 1.25rem 0; overflow: hidden; }
    .ats-score-fill { height: 100%; border-radius: 6px; transition: width 0.8s cubic-bezier(0.4,0,0.2,1); }
    .ats-score-label { font-size: 1.8rem; font-weight: 800; letter-spacing: -0.02em; margin-right: 0.3rem; }
    .ats-tag-group { display: flex; flex-wrap: wrap; gap: 0.45rem; margin: 0.5rem 0 0.85rem 0; }
    .ats-tag { display: inline-flex; align-items: center; gap: 0.25rem; padding: 0.3rem 0.7rem;
        border-radius: 20px; font-size: 0.82rem; font-weight: 500; line-height: 1.4; }
    .ats-tag-matched { background: rgba(34,197,94,0.10); color: #15803d; border: 1px solid rgba(34,197,94,0.20); }
    .ats-tag-missing { background: rgba(245,158,11,0.10); color: #b45309; border: 1px solid rgba(245,158,11,0.20); }
    .ats-tag-suggested { background: rgba(99,102,241,0.08); color: #4f46e5; border: 1px solid rgba(99,102,241,0.18); }
    .ats-section-label { font-weight: 700; font-size: 0.88rem; margin-bottom: 0.35rem; }

    /* 白色玻璃表面（共用 backdrop + 背景） */
    .stExpander, .stAlert, [data-testid="stSidebar"], [data-testid="stToast"] {
        background: rgba(255, 255, 255, 0.88);
        backdrop-filter: blur(16px); -webkit-backdrop-filter: blur(16px);
        border: 1px solid rgba(148, 163, 184, 0.25);
    }
    .stExpander { border-radius: 16px !important; box-shadow: 0 1px 3px rgba(0,0,0,0.04); }
    .stExpander:hover { border-color: rgba(99, 102, 241, 0.25) !important; }
    .stExpander p, .stExpander span, .stExpander label { color: #1e293b !important; }
    .stAlert { border-radius: 14px; color: #1e293b !important; }
    [data-testid="stSidebar"] { border-right: 1px solid rgba(148, 163, 184, 0.25); }
    [data-testid="stSidebar"] * { color: #1e293b !important; }
    [data-testid="stToast"] { background: rgba(255, 255, 255, 0.92) !important; border-radius: 14px; color: #1e293b; }

    /* 滚动条 */
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb { background: rgba(0,0,0,0.15); border-radius: 3px; }
    ::-webkit-scrollbar-thumb:hover { background: rgba(0,0,0,0.25); }

    /* 全局文字颜色 */
    html, body, .stApp, .stText { color: #1e293b !important; }
    .stMarkdown p, .stMarkdown span, .stMarkdown li,
    [data-testid="stMarkdownContainer"] p, [data-testid="stMarkdownContainer"] li, [data-testid="stMarkdownContainer"] span,
    .stAlert p, .stAlert span { color: #1e293b; }
    h1, h2, h3, h4, h5, h6 { color: #1e293b !important; }
    label, caption, .stCaption, .stSpinner { color: #475569 !important; }
    input, textarea, select { color: #1e293b !important; }
    .module-card-body p, .module-card-body li, .module-card-body span { color: #1e293b !important; }
    a { color: #4f46e5; text-decoration: none; }
    a:hover { color: #6366f1; }
    strong { color: #0f172a; font-weight: 600; }
</style>
""", unsafe_allow_html=True)

# ==================== 页面标题（无卡片背景）====================
st.markdown('<h1 class="hero-title">AI 简历优化助手</h1>', unsafe_allow_html=True)
st.markdown('<div class="hero-accent"></div>', unsafe_allow_html=True)
st.markdown(
    '<p class="hero-subtitle">上传简历，输入岗位要求，AI 深度分析匹配度并生成专业优化建议</p>',
    unsafe_allow_html=True,
)

# ==================== 文件解析函数 ====================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """从 PDF 文件的字节数据中提取文字"""
    pdf_reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 DOCX 文件的字节数据中提取文字"""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    return "\n".join(text_parts)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """从 TXT 文件的字节数据中提取文字"""
    # 尝试 UTF-8 编码，失败则回退到 GBK（兼容中文 Windows 系统）
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("gbk", errors="ignore")


def extract_resume_text(uploaded_file) -> str | None:
    """
    根据文件类型选择对应的提取方法，返回提取的文字。
    如果文件类型不支持，返回 None。
    """
    file_bytes = uploaded_file.read()  # 读取文件的原始字节
    file_type = uploaded_file.name.lower()  # 获取文件名并转小写

    if file_type.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif file_type.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif file_type.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        return None


def extract_resume_from_bytes(file_bytes: bytes, file_name: str) -> str | None:
    """从已读取的字节数据中提取简历文字（用于 session_state 缓存场景）。"""
    file_name_lower = file_name.lower()
    if file_name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif file_name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif file_name_lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    return None


# ==================== DeepSeek API 配置 ====================
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)


def build_analysis_prompt(resume: str, job_desc: str) -> str:
    """构建 AI 分析提示词，7 个模块 + 1 个 ATS 分析，用分隔标记分隔。"""
    prompt = f"""你是一位拥有 15 年经验的资深 HR 总监兼 Career Coach，曾任职于顶尖科技公司和猎头机构。你的专长是简历优化、岗位匹配度分析和职业发展规划。

请严格按照以下规范，分析候选人简历与目标岗位的匹配情况。

【岗位描述】
{job_desc}

【候选人简历】
{resume}

【核心规则】
1. 绝对禁止编造。不得虚构任何公司名、学校名、项目名、实习经历、证书、技能、量化数据、成果数字。
2. 只能基于候选人原始简历中「已有」的信息进行优化表达和结构调整。
3. 如果原始简历缺少某项关键信息，不得靠猜测填充。必须在第 7 模块中列出。
4. 输出的正文用中文撰写，但以下内容保持英文原样：技术名词（Python, SQL, R, Excel）、专业术语（Machine Learning, Data Analysis, Biostatistics）、公司名称、产品名称、行业标准缩写。
5. 岗位描述中的英文关键词也保持原样。
6. 每个模块的输出必须紧跟在其 ===SECTION_N=== 标记之后，不要在其他位置出现这些标记。

【评分标准】
- 85-100：高度匹配，核心技能和经验与岗位要求高度吻合
- 70-84：较好匹配，主要条件满足，少数次要条件欠缺
- 50-69：部分匹配，有一定基础但存在明显差距
- 0-49：匹配度低，背景与岗位要求差距较大
评分时需综合考量：硬技能、行业经验、学历背景、项目经验、软技能 5 个维度。

【输出格式 —— 7 个模块 + 1 个 ATS 分析，严格按顺序】

===SECTION_1===
## 匹配度评分
（格式：{{分数}}/100，单独一行。然后另起一段，从硬技能、行业经验、学历背景、项目经验、软技能 5 个维度简要说明评分依据。最后用一句话总结该候选人的核心竞争力。）

===SECTION_2===
## 岗位关键词匹配
（从岗位描述中提取 8-15 个核心关键词，逐一分析。对每个关键词标注 ✅已覆盖 或 ⚠️未覆盖。未覆盖的关键词需给出具体建议，说明应在简历哪个部分补充。）

===SECTION_ATS===
## ATS 关键词分析
（这是专门的 ATS（Applicant Tracking System）关键词扫描模块。请严格按照以下格式输出，方便程序解析后做可视化展示。）

**ATS_SCORE:{{分数}}**

### ✅ 已匹配
（列出已在简历中找到的关键词，每行一个。格式：`- 关键词（所在部分）`）

### ⚠️ 缺失
（列出岗位要求但简历中未找到的关键词，每行一个。格式：`- 关键词 → 建议在「XX部分」自然加入`）

### ⭐ 建议补充
（列出可以进一步提升匹配度的补充关键词，每行一个。格式：`- 关键词 → 建议在「XX部分」体现`）

（注意：已匹配列表中已有的关键词，不要在缺失或建议补充中重复列出。只基于简历已有信息判断，禁止编造。）

===SECTION_3===
## 简历主要问题
（以 - 列表形式，逐条指出简历与岗位要求之间的主要差距。每条问题包含：问题描述 → 为什么重要 → 具体改进方向。按严重程度排序，最重要的问题排前面。）

===SECTION_4===
## 可量化优化建议
（以 - 列表形式，给出可直接执行的具体优化建议。每条建议包含：优化目标、具体操作步骤、预期效果。优先给出能将主观描述转化为量化成果的建议。如果原始简历缺少量化数据，写「[建议补充：具体数据/成果]」而不是编造数字。）

===SECTION_5===
## 优化后的 Bullet Points
（以 - 列表形式，输出 6-12 条优化后的简历要点。每条 bullet point 遵循 STAR 法则或「动作动词 + 具体内容 + 量化成果」结构。内容必须来源于候选人真实经历，关键词需贴近岗位描述。每条 bullet 控制在 1-2 行，简洁有力。）

===SECTION_6===
## 优化后的完整简历
（基于原始简历，输出一版完整的优化后中文简历。保持原有结构但优化表达。信息不足处用「[建议补充：需要补充的内容说明]」标记。英文术语和公司名保持原样。简历应包含：个人信息区、求职意向、教育背景、工作/实习经历、项目经历、技能、证书（如有）。）

===SECTION_7===
## 需要补充的信息
（以 - 列表形式，列出原始简历中缺失但岗位要求的全部关键信息。每条注明：缺失项、对岗位的重要性、补充建议。如果简历信息已足够完整，写「简历信息较完整，暂无需补充的关键信息。」）
"""
    return prompt


def call_deepseek_api(prompt: str) -> str | None:
    """
    调用 DeepSeek API 进行分析。
    成功时返回 AI 回复内容，失败时返回 None 并在页面显示错误。
    """
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是一位拥有15年经验的资深HR总监兼职业教练，专长是简历优化、岗位匹配分析和职业发展规划。你的分析专业、具体、可执行，从不编造信息。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=8192,
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"❌ AI 分析失败：{e}")
        return None


def parse_ai_sections(ai_output: str) -> dict[str, str]:
    """将 AI 输出按 ===SECTION_N=== / ===SECTION_ATS=== 拆分为 8 个模块。"""
    section_map = {
        "===SECTION_1===": "score",
        "===SECTION_2===": "keyword_match",
        "===SECTION_ATS===": "ats_analysis",
        "===SECTION_3===": "problems",
        "===SECTION_4===": "quantified_suggestions",
        "===SECTION_5===": "bullets",
        "===SECTION_6===": "full_resume",
        "===SECTION_7===": "missing_info",
    }
    sections = {v: "" for v in section_map.values()}
    parts = {}
    current_key = None
    for line in ai_output.split("\n"):
        stripped = line.strip()
        if stripped in section_map:
            current_key = section_map[stripped]
            continue
        if current_key:
            parts[current_key] = parts.get(current_key, "") + line + "\n"
    for key in sections:
        sections[key] = parts.get(key, "").strip()
        if not sections[key]:
            sections[key] = "（此模块解析失败，请重试）"
    return sections


def parse_ats_section(ats_text: str) -> dict:
    """解析 ATS 分析文本，提取分数和分类关键词列表。"""
    result = {"score": 0, "matched": [], "missing": [], "suggested": []}
    if not ats_text or "解析失败" in ats_text:
        return result
    # 提取分数
    score_match = re.search(r"ATS_SCORE[:\s]*(\d+)", ats_text)
    if score_match:
        result["score"] = int(score_match.group(1))
    # 分割三大区域
    for _section, marker, result_key in [
        ("已匹配", "### ✅", "matched"),
        ("缺失", "### ⚠️", "missing"),
        ("建议补充", "### ⭐", "suggested"),
    ]:
        pattern = re.escape(marker) + r"\s*[^\n]*\n(.*?)(?=###\s|$)"
        match = re.search(pattern, ats_text, re.DOTALL)
        if match:
            for line in match.group(1).strip().split("\n"):
                line = line.strip().lstrip("- ").strip()
                if not line:
                    continue
                # 尝试拆分关键词和位置/建议
                if "→" in line:
                    kw, loc = line.split("→", 1)
                    result[result_key].append((kw.strip(), loc.strip()))
                elif "（" in line and "）" in line:
                    # 格式：关键词（位置）
                    kw_start = line.rfind("（")
                    kw = line[:kw_start].strip()
                    loc = line[kw_start:].strip("（）")
                    result[result_key].append((kw, loc))
                else:
                    result[result_key].append((line, ""))
    return result


def generate_word_report(sections: dict[str, str]) -> bytes:
    """生成 Word 文档：标题 + 8 个分析模块。"""
    doc = Document()
    title = doc.add_heading("AI简历优化结果", level=0)
    for run in title.runs:
        run.font.size = Pt(22)
    for heading, key in [
        ("匹配度评分", "score"), ("岗位关键词匹配", "keyword_match"),
        ("ATS 关键词分析", "ats_analysis"), ("简历主要问题", "problems"),
        ("可量化优化建议", "quantified_suggestions"), ("优化后的 Bullet Points", "bullets"),
        ("优化后的完整简历", "full_resume"), ("需要补充的信息", "missing_info"),
    ]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(sections[key])
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


# ==================== 输入区域（白色玻璃主面板）====================
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
col1, col2 = st.columns(2)

with col1:
    st.markdown(
        '<p style="font-weight:600;font-size:0.92rem;color:#475569;margin-bottom:0.35rem;">'
        '<span style="font-size:1.1rem;">📂</span> 上传简历文件</p>',
        unsafe_allow_html=True,
    )
    uploaded_file = st.file_uploader(
        label="上传简历文件",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=False,
        help="支持 PDF、DOCX、TXT 格式。上传后优先使用文件内容。",
        label_visibility="collapsed",
    )

    # 即时验证上传文件，缓存内容到 session_state
    if uploaded_file is not None:
        cache_key = f"{uploaded_file.name}_{uploaded_file.size}"
        if st.session_state.get("_resume_cache_key") != cache_key:
            file_bytes = uploaded_file.read()
            file_name_lower = uploaded_file.name.lower()
            if not (file_name_lower.endswith(".pdf") or file_name_lower.endswith(".docx") or file_name_lower.endswith(".txt")):
                st.error("❌ 不支持的文件格式，请上传 PDF、DOCX 或 TXT 文件。")
                st.session_state._resume_cache_key = None
                st.session_state._cached_resume_text = None
            else:
                extracted = extract_resume_from_bytes(file_bytes, uploaded_file.name)
                if extracted is None or not extracted.strip():
                    st.error("❌ 未能从文件中提取到文字，请检查文件内容是否为空或损坏。")
                    st.session_state._resume_cache_key = None
                    st.session_state._cached_resume_text = None
                else:
                    st.success("✅ 文件上传成功，已读取简历内容。")
                    st.session_state._resume_cache_key = cache_key
                    st.session_state._cached_resume_text = extracted
    else:
        st.session_state.pop("_resume_cache_key", None)
        st.session_state.pop("_cached_resume_text", None)

    st.markdown(
        '<p style="text-align:center;color:#64748b;margin:0.65rem 0;font-size:0.85rem;">— 或者 —</p>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p style="font-weight:600;font-size:0.92rem;color:#475569;margin-bottom:0.35rem;">'
        '<span style="font-size:1.1rem;">📝</span> 直接粘贴简历内容</p>',
        unsafe_allow_html=True,
    )
    resume_text = st.text_area(
        label="直接粘贴简历内容",
        placeholder="在这里粘贴你的简历文字...",
        height=160,
        label_visibility="collapsed",
    )

with col2:
    st.markdown(
        '<p style="font-weight:600;font-size:0.92rem;color:#475569;margin-bottom:0.35rem;">'
        '<span style="font-size:1.1rem;">💼</span> 岗位描述</p>',
        unsafe_allow_html=True,
    )
    job_text = st.text_area(
        label="岗位描述",
        placeholder="在此粘贴目标岗位的完整 JD（职位描述），信息越详细分析越精准...",
        height=390,
        label_visibility="collapsed",
    )
st.markdown('</div>', unsafe_allow_html=True)

# ==================== 分析按钮（居中）====================
_, btn_col, _ = st.columns([1, 1, 1])
with btn_col:
    analyze_clicked = st.button("✨ 开始 AI 分析", use_container_width=True, key="analyze_btn")

# ==================== 点击按钮后的处理逻辑 ====================
if analyze_clicked:
    # --- 第一步：判断是否有简历内容（文件或手动输入）---
    has_file = uploaded_file is not None          # 用户是否上传了文件
    has_text = bool(resume_text.strip())           # 用户是否手动粘贴了文字

    if not has_file and not has_text:
        st.warning("⚠️ 请上传简历文件，或在输入框中粘贴简历内容。")
    # --- 第二步：检查是否填写了岗位描述 ---
    elif not job_text.strip():
        st.warning("⚠️ 请填写「岗位描述」后再进行分析。")
    else:
        # --- 第三步：优先使用文件内容，否则使用手动输入 ---
        if has_file:
            cached = st.session_state.get("_cached_resume_text")
            if cached:
                final_resume = cached
                source_hint = "（从文件中提取）"
            else:
                st.stop()
        else:
            # 用户没有上传文件，使用手动输入的简历文字
            final_resume = resume_text.strip()
            source_hint = "（手动输入）"

        # --- 第四步：检查 API Key 是否已配置 ---
        if not DEEPSEEK_API_KEY:
            st.error("❌ 未配置 DeepSeek API Key")
            st.info("请在项目根目录创建 `.env` 文件，写入：\n\n```\nDEEPSEEK_API_KEY=你的key\n```\n\n然后重启 Streamlit。")
            st.stop()

        # 输入读取成功的提示
        st.toast("✅ 简历读取成功，正在启动 AI 分析...", icon="✅")

        # 使用折叠面板展示输入内容
        with st.expander("📋 查看输入内容", expanded=False):
            preview_col1, preview_col2 = st.columns(2)
            with preview_col1:
                st.markdown(f"**📝 简历内容** {source_hint}")
                st.text(final_resume)
            with preview_col2:
                st.markdown("**💼 岗位描述**")
                st.text(job_text)

        st.divider()

        # --- 第五步：调用 AI 进行分析 ---
        with st.spinner("🤖 AI 正在深度分析你的简历与岗位匹配度..."):
            prompt = build_analysis_prompt(final_resume, job_text)
            ai_result = call_deepseek_api(prompt)

        # --- 第六步：解析并展示 AI 结果（卡片式布局）---
        if ai_result:
            sections = parse_ai_sections(ai_result)

            st.markdown('<p class="results-heading">🤖 AI 分析报告</p>', unsafe_allow_html=True)

            # --- 前两个模块 ---
            for icon, title, key in [
                ("📊", "匹配度评分", "score"),
                ("🎯", "岗位关键词匹配", "keyword_match"),
            ]:
                st.markdown(f"""<div class="module-card">
                <div class="module-card-title"><span class="module-card-icon">{icon}</span> {title}</div>
                <div class="module-card-body">{sections[key]}</div>
                </div>""", unsafe_allow_html=True)

            # --- ATS 关键词分析（特殊卡片） ---
            ats_data = parse_ats_section(sections.get("ats_analysis", ""))
            ats_score = ats_data.get("score", 0)
            # 颜色：<40 红，40-69 黄，70-84 蓝绿，85+ 绿
            if ats_score >= 85:
                bar_color = "#22c55e"
            elif ats_score >= 70:
                bar_color = "#06b6d4"
            elif ats_score >= 40:
                bar_color = "#f59e0b"
            else:
                bar_color = "#ef4444"

            st.markdown(f"""<div class="module-card">
            <div class="module-card-title"><span class="module-card-icon">🤖</span> ATS 关键词分析</div>
            <div class="module-card-body">
            <div style="display:flex;align-items:baseline;gap:0.5rem;">
                <span class="ats-score-label" style="color:{bar_color};">{ats_score}</span>
                <span style="color:#475569;font-size:0.95rem;">/ 100 匹配率</span>
            </div>
            <div class="ats-score-bar"><div class="ats-score-fill" style="width:{ats_score}%;background:{bar_color};"></div></div>
            """, unsafe_allow_html=True)

            for label, items, css_class in [
                ("✅ 已匹配关键词", ats_data.get("matched", []), "ats-tag-matched"),
                ("⚠️ 缺失关键词", ats_data.get("missing", []), "ats-tag-missing"),
                ("⭐ 建议补充关键词", ats_data.get("suggested", []), "ats-tag-suggested"),
            ]:
                if items:
                    st.markdown(f'<p class="ats-section-label">{label} <span style="color:#64748b;font-weight:400;">({len(items)})</span></p>', unsafe_allow_html=True)
                    tags_html = "".join(
                        f'<span class="ats-tag {css_class}">{kw}<span style="font-size:0.72rem;opacity:0.7;">{" — " + loc if loc else ""}</span></span>'
                        for kw, loc in items
                    )
                    st.markdown(f'<div class="ats-tag-group">{tags_html}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<p class="ats-section-label">{label} <span style="color:#64748b;font-weight:400;">(0)</span></p>', unsafe_allow_html=True)
                    st.markdown('<p style="color:#64748b;font-size:0.85rem;margin-bottom:0.5rem;">暂无数据</p>', unsafe_allow_html=True)

            st.markdown('</div></div>', unsafe_allow_html=True)

            # --- 后五个模块 ---
            for icon, title, key in [
                ("🔍", "简历主要问题", "problems"),
                ("📈", "可量化优化建议", "quantified_suggestions"),
                ("✨", "优化后的 Bullet Points", "bullets"),
                ("📝", "优化后的完整简历", "full_resume"),
                ("📋", "需要补充的信息", "missing_info"),
            ]:
                st.markdown(f"""<div class="module-card">
                <div class="module-card-title"><span class="module-card-icon">{icon}</span> {title}</div>
                <div class="module-card-body">{sections[key]}</div>
                </div>""", unsafe_allow_html=True)

            # 操作区域：复制 + 下载
            st.markdown('<div class="module-card">', unsafe_allow_html=True)
            st.markdown(
                '<p style="font-weight:700;color:#1e293b;margin-bottom:0.75rem;">📋 导出与复制</p>',
                unsafe_allow_html=True,
            )
            st.text_area(
                label="优化后完整简历（点击全选后复制）",
                value=sections["full_resume"],
                height=220,
                key="copy_area",
                label_visibility="collapsed",
            )
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                word_bytes = generate_word_report(sections)
                st.download_button(
                    label="📥 下载 Word 文档",
                    data=word_bytes,
                    file_name="优化后简历.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            st.markdown('</div>', unsafe_allow_html=True)

